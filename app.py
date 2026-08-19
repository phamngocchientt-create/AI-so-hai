import io
import json
import streamlit as st
import requests
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ==========================================
# CẤU HÌNH TRANG WEB STREAMLIT
# ==========================================
st.set_page_config(
    page_title="EduPlan AI - Trợ Lý Soạn Giáo Án",
    page_icon="📚",
    layout="wide",
    initial_sidebar_state="expanded"
)

st.markdown("""
    <style>
    .main-title { font-size: 28px; font-weight: bold; color: #1E3A8A; margin-bottom: 5px; }
    .sub-title { font-size: 15px; color: #4B5563; margin-bottom: 20px; }
    .stButton>button { width: 100%; border-radius: 8px; font-weight: bold; background-color: #2563EB; color: white; }
    </style>
""", unsafe_allow_html=True)

if "lesson_content" not in st.session_state:
    st.session_state["lesson_content"] = ""
if "lesson_title" not in st.session_state:
    st.session_state["lesson_title"] = "Ke_hoach_bai_day"

# ==========================================
# CÁC HÀM XỬ LÝ FILE VÀ GỌI GEMINI API TRỰC TIẾP
# ==========================================
def extract_text_from_docx(file_bytes):
    doc = Document(file_bytes)
    full_text = []
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text.strip())
    for table in doc.tables:
        full_text.append("\n[BẢNG MẪU HOẠT ĐỘNG]")
        for row in table.rows:
            row_data = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            full_text.append(" | ".join(row_data))
    return "\n".join(full_text)

def generate_with_gemini_api(api_key, system_instruction, prompt_text):
    """Gọi trực tiếp Google Gemini API qua HTTP POST request"""
    url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-flash:generateContent?key={api_key}"
    headers = {"Content-Type": "application/json"}
    payload = {
        "contents": [{"parts": [{"text": prompt_text}]}],
        "systemInstruction": {"parts": [{"text": system_instruction}]},
        "generationConfig": {"temperature": 0.4}
    }
    
    response = requests.post(url, headers=headers, json=payload, timeout=60)
    if response.status_code != 200:
        error_msg = response.json().get("error", {}).get("message", response.text)
        raise Exception(f"Lỗi API Google ({response.status_code}): {error_msg}")
    
    data = response.json()
    candidates = data.get("candidates", [])
    if candidates:
        return candidates[0]["content"]["parts"][0]["text"]
    return "Không nhận được phản hồi từ AI."

def generate_lesson_plan(api_key, mode, subject, grade, textbook, title, duration, options, sample_content=""):
    options_text = ", ".join(options) if options else "Theo chuẩn sư phạm hiện hành"
    
    if mode == "5512":
        sys_ins = (
            "Bạn là chuyên gia sư phạm hàng đầu tại Việt Nam. "
            "Hãy soạn kế hoạch bài dạy (giáo án) chuẩn xác theo định dạng CÔNG VĂN 5512 của Bộ Giáo dục & Đào tạo. "
            "Cấu trúc bài dạy:\n"
            "TÊN BÀI DẠY (Môn, Lớp, Bộ sách, Thời lượng)\n"
            "I. MỤC TIÊU (1. Kiến thức, 2. Năng lực chung & đặc thù, 3. Phẩm chất)\n"
            "II. THIẾT BỊ DẠY HỌC VÀ HỌC LIỆU (1. Giáo viên, 2. Học sinh)\n"
            "III. TIẾN TRÌNH DẠY HỌC (4 hoạt động: 1. Khởi động, 2. Hình thành kiến thức mới, 3. Luyện tập, 4. Vận dụng). "
            "Mỗi hoạt động có: a. Mục tiêu, b. Nội dung, c. Sản phẩm, d. Tổ chức thực hiện (Bước 1: Chuyển giao, Bước 2: Thực hiện, Bước 3: Báo cáo thảo luận, Bước 4: Kết luận).\n"
            "IV. PHỤ LỤC (Nếu có)"
        )
        prompt = (
            f"Hãy soạn giáo án chi tiết:\n"
            f"- Môn học: {subject}\n"
            f"- Khối lớp: {grade}\n"
            f"- Bộ sách: {textbook}\n"
            f"- Tên bài: {title}\n"
            f"- Thời lượng: {duration}\n"
            f"- Yêu cầu bổ trợ: {options_text}\n"
        )
    else:
        sys_ins = (
            "Bạn là trợ lý soạn giáo án thông minh. "
            "Hãy học và làm theo cấu trúc, cách đặt đề mục, bảng biểu và văn phong từ bài giáo án mẫu để viết bài mới."
        )
        prompt = (
            f"Dưới đây là GIÁO ÁN MẪU:\n"
            f"--- BẮT ĐẦU MẪU ---\n{sample_content[:3500]}\n--- KẾT THÚC MẪU ---\n\n"
            f"Thông tin bài mới:\n"
            f"- Môn: {subject} - Khối: {grade} - Bộ sách: {textbook}\n"
            f"- Tên bài mới: {title} ({duration})\n"
            f"- Yêu cầu: {options_text}\n"
        )
    return generate_with_gemini_api(api_key, sys_ins, prompt)

def create_docx_file(content, title):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.79)
        section.bottom_margin = Inches(0.79)
        section.left_margin = Inches(1.18)
        section.right_margin = Inches(0.59)

    for line in content.split("\n"):
        line = line.strip()
        if not line:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(3)
        
        if any(line.startswith(prefix) for prefix in ["I.", "II.", "III.", "IV.", "V."]):
            run = p.add_run(line)
            run.font.name = "Times New Roman"
            run.font.size = Pt(13)
            run.font.bold = True
            run.font.color.rgb = RGBColor(30, 58, 138)
        elif line.startswith("#") or line.startswith("KẾ HOẠCH BÀI DẠY") or line.startswith("TÊN BÀI DẠY"):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line.replace("#", "").strip())
            run.font.name = "Times New Roman"
            run.font.size = Pt(14)
            run.font.bold = True
        elif any(line.startswith(prefix) for prefix in ["1.", "2.", "3.", "4.", "a.", "b.", "c.", "d.", "Hoạt động"]):
            run = p.add_run(line)
            run.font.name = "Times New Roman"
            run.font.size = Pt(13)
            run.font.bold = True
        else:
            run = p.add_run(line)
            run.font.name = "Times New Roman"
            run.font.size = Pt(13)

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# ==========================================
# GIAO DIỆN CHÍNH
# ==========================================
st.markdown("<div class='main-title'>📚 EduPlan AI - Trợ Lý Soạn Giáo Án</div>", unsafe_allow_html=True)
st.markdown("<div class='sub-title'>Tự động soạn kế hoạch bài dạy chuẩn Công văn 5512 hoặc học theo mẫu Word của trường bạn</div>", unsafe_allow_html=True)

col_left, col_right = st.columns([1.1, 1.9], gap="medium")

with col_left:
    st.subheader("⚙️ Thiết lập bài soạn")
    
    default_api_key = st.secrets.get("GEMINI_API_KEY", "") if "GEMINI_API_KEY" in st.secrets else ""
    gemini_api_key = st.text_input("🔑 Google Gemini API Key:", value=default_api_key, type="password", help="Lấy tại aistudio.google.com")
    
    tab1, tab2 = st.tabs(["⚡ Chuẩn Công văn 5512", "📁 Dựa theo Mẫu Word"])
    
    with tab1:
        st.caption("Sinh giáo án chuẩn mực 4 hoạt động của Bộ GD&ĐT")
        c1, c2 = st.columns(2)
        with c1:
            subject_5512 = st.selectbox("Môn học:", ["Khoa học tự nhiên (Hóa học)", "Khoa học tự nhiên (Sinh học)", "Khoa học tự nhiên (Vật lí)", "Toán học", "Ngữ văn", "Lịch sử & Địa lí", "Khác"])
            grade_5512 = st.selectbox("Khối lớp:", ["Lớp 6", "Lớp 7", "Lớp 8", "Lớp 9"])
        with c2:
            book_5512 = st.selectbox("Bộ sách:", ["Cánh Diều", "Kết Nối Tri Thức", "Chân Trời Sáng Tạo"])
            duration_5512 = st.selectbox("Thời lượng:", ["1 tiết", "2 tiết", "3 tiết", "4 tiết"])
            
        title_5512 = st.text_input("Tên bài học:", value="Bài: Axit - Tính chất và ứng dụng")
        
        st.write("Tùy chọn bổ trợ:")
        opt_ws = st.checkbox("Kèm Phiếu học tập (Worksheet)", value=True)
        opt_quiz = st.checkbox("Kèm 5 câu trắc nghiệm củng cố", value=True)
        opt_game = st.checkbox("Thiết kế trò chơi khởi động / Thí nghiệm", value=True)
        
        options_list_5512 = []
        if opt_ws: options_list_5512.append("Kèm phiếu học tập chi tiết")
        if opt_quiz: options_list_5512.append("Kèm 5 câu hỏi trắc nghiệm đánh giá")
        if opt_game: options_list_5512.append("Thiết kế trò chơi/thí nghiệm khởi động sinh động")
        
        btn_generate_5512 = st.button("🚀 BẮT ĐẦU SOẠN (CHUẨN 5512)", key="btn_5512")

    with tab2:
        st.caption("AI sẽ đọc cấu trúc bảng biểu và phong cách của file mẫu để viết bài mới")
        uploaded_file = st.file_uploader("Tải lên file giáo án mẫu (.docx):", type=["docx"])
        
        c1, c2 = st.columns(2)
        with c1:
            subject_custom = st.selectbox("Môn bài mới:", ["Khoa học tự nhiên (Hóa học)", "Khoa học tự nhiên (Sinh học)", "Toán học", "Ngữ văn", "Khác"], key="sb_custom")
            grade_custom = st.selectbox("Khối lớp mới:", ["Lớp 6", "Lớp 7", "Lớp 8", "Lớp 9"], key="gr_custom")
        with c2:
            duration_custom = st.selectbox("Thời lượng bài mới:", ["1 tiết", "2 tiết", "3 tiết"], key="dur_custom")
            book_custom = st.text_input("Bộ sách:", value="Cánh Diều", key="bk_custom")
            
        title_custom = st.text_input("Tên bài học mới cần soạn:", value="Bài: Bazơ - Tính chất hóa học", key="tit_custom")
        
        btn_generate_custom = st.button("🚀 SOẠN BÀI THEO MẪU ĐÃ TẢI", key="btn_custom")

if btn_generate_5512:
    if not gemini_api_key:
        st.error("⚠️ Vui lòng nhập Google Gemini API Key để tiếp tục!")
    elif not title_5512:
        st.warning("⚠️ Vui lòng nhập tên bài học!")
    else:
        with st.spinner("⏳ Đang soạn giáo án theo chuẩn 5512..."):
            try:
                res = generate_lesson_plan(
                    api_key=gemini_api_key,
                    mode="5512",
                    subject=subject_5512,
                    grade=grade_5512,
                    textbook=book_5512,
                    title=title_5512,
                    duration=duration_5512,
                    options=options_list_5512
                )
                st.session_state["lesson_content"] = res
                st.session_state["lesson_title"] = f"GiaoAn_{grade_5512}_{title_5512.replace(' ', '_')}"
                st.success("✅ Đã tạo xong giáo án thành công!")
            except Exception as e:
                st.error(f"❌ Có lỗi xảy ra: {str(e)}")

if btn_generate_custom:
    if not gemini_api_key:
        st.error("⚠️ Vui lòng nhập Google Gemini API Key để tiếp tục!")
    elif not uploaded_file:
        st.warning("⚠️ Vui lòng tải lên file Word giáo án mẫu (.docx)!")
    elif not title_custom:
        st.warning("⚠️ Vui lòng nhập tên bài học mới cần soạn!")
    else:
        with st.spinner("⏳ Đang đọc mẫu và soạn bài mới..."):
            try:
                sample_text = extract_text_from_docx(uploaded_file)
                res = generate_lesson_plan(
                    api_key=gemini_api_key,
                    mode="custom",
                    subject=subject_custom,
                    grade=grade_custom,
                    textbook=book_custom,
                    title=title_custom,
                    duration=duration_custom,
                    options=["Bắt chước phong cách bài mẫu"],
                    sample_content=sample_text
                )
                st.session_state["lesson_content"] = res
                st.session_state["lesson_title"] = f"GiaoAn_TheoMau_{title_custom.replace(' ', '_')}"
                st.success("✅ Đã hoàn thành bài soạn theo mẫu!")
            except Exception as e:
                st.error(f"❌ Có lỗi xảy ra: {str(e)}")

with col_right:
    st.subheader("📄 Bản xem trước & Chỉnh sửa trực tiếp")
    
    if st.session_state["lesson_content"]:
        edited_text = st.text_area(
            "Gõ sửa trực tiếp trước khi tải về:",
            value=st.session_state["lesson_content"],
            height=550
        )
        st.session_state["lesson_content"] = edited_text
        
        docx_file = create_docx_file(st.session_state["lesson_content"], st.session_state["lesson_title"])
        
        st.download_button(
            label="📥 TẢI XUỐNG FILE WORD (.DOCX) CHUẨN IN ẤN",
            data=docx_file,
            file_name=f"{st.session_state['lesson_title']}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary"
        )
    else:
        st.info("👈 Chọn thông tin ở cột bên trái và bấm nút **Bắt đầu soạn** để xem kết quả tại đây.")
