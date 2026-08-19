import io
import streamlit as st
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from google import genai
from google.genai import types

# ==========================================
# CẤU HÌNH TRANG WEB STREAMLIT
# ==========================================
st.set_page_config(
    page_title="EduPlan AI - Trợ Lý Soạn Giáo Án",
    page_icon="📚",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Tùy chỉnh CSS để giao diện trực quan, rõ ràng
st.markdown("""
    <style>
    .main-title { font-size: 28px; font-weight: bold; color: #1E3A8A; margin-bottom: 5px; }
    .sub-title { font-size: 15px; color: #4B5563; margin-bottom: 20px; }
    .stButton>button { width: 100%; border-radius: 8px; font-weight: bold; background-color: #2563EB; color: white; }
    </style>
""", unsafe_allow_html=True)

# Khởi tạo session state lưu nội dung giáo án
if "lesson_content" not in st.session_state:
    st.session_state["lesson_content"] = ""
if "lesson_title" not in st.session_state:
    st.session_state["lesson_title"] = "Ke_hoach_bai_day"

# ==========================================
# CÁC HÀM HỖ TRỢ XỬ LÝ FILE VÀ AI
# ==========================================
def extract_text_from_docx(file_bytes):
    """Trích xuất toàn bộ văn bản và bảng biểu từ file Word giáo án mẫu"""
    doc = Document(file_bytes)
    full_text = []
    
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text.strip())
            
    for table in doc.tables:
        full_text.append("\n[BẢNG MẪU HOẠT ĐỘNG]")
        for row in table.rows:
            row_data = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            full_text.append(" | ".join(row_data))
            
    return "\n".join(full_text)

def generate_lesson_plan(api_key, mode, subject, grade, textbook, title, duration, options, sample_content=""):
    """Gọi Gemini API để sinh nội dung giáo án chuẩn"""
    client = genai.Client(api_key=api_key)
    
    options_text = ", ".join(options) if options else "Theo chuẩn sư phạm hiện hành"
    
    if mode == "5512":
        system_instruction = (
            "Bạn là chuyên gia sư phạm hàng đầu tại Việt Nam. "
            "Hãy soạn kế hoạch bài dạy (giáo án) chuẩn xác theo định dạng CÔNG VĂN 5512 của Bộ Giáo dục & Đào tạo. "
            "Cấu trúc gồm:\n"
            "TÊN BÀI DẠY (Môn, Lớp, Bộ sách, Thời lượng)\n"
            "I. MỤC TIÊU (1. Kiến thức, 2. Năng lực chung & đặc thù, 3. Phẩm chất)\n"
            "II. THIẾT BỊ DẠY HỌC VÀ HỌC LIỆU (1. Giáo viên, 2. Học sinh)\n"
            "III. TIẾN TRÌNH DẠY HỌC (Gồm 4 hoạt động: 1. Khởi động, 2. Hình thành kiến thức mới, 3. Luyện tập, 4. Vận dụng). "
            "Mỗi hoạt động phải trình bày rõ 4 mục: a. Mục tiêu, b. Nội dung, c. Sản phẩm, d. Tổ chức thực hiện (Bước 1: Chuyển giao nhiệm vụ, Bước 2: Thực hiện nhiệm vụ, Bước 3: Báo cáo thảo luận, Bước 4: Kết luận nhận định).\n"
            "IV. PHỤ LỤC (Nếu có yêu cầu phiếu học tập/câu hỏi củng cố)."
        )
        prompt = (
            f"Hãy soạn giáo án chi tiết cho bài học sau:\n"
            f"- Môn học: {subject}\n"
            f"- Khối lớp: {grade}\n"
            f"- Bộ sách: {textbook}\n"
            f"- Tên bài học: {title}\n"
            f"- Thời lượng: {duration}\n"
            f"- Yêu cầu bổ trợ: {options_text}\n"
            f"Văn phong sư phạm chuẩn mực, hoạt động chi tiết, dễ dàng triển khai trên lớp."
        )
    else:
        system_instruction = (
            "Bạn là trợ lý soạn giáo án thông minh. "
            "Nhiệm vụ của bạn là học và 'bắt chước' 100% cấu trúc, cách đặt đề mục, cách phân chia hoạt động và văn phong sư phạm "
            "từ bài giáo án mẫu được cung cấp để viết cho một bài học mới."
        )
        prompt = (
            f"Dưới đây là nội dung GIÁO ÁN MẪU của giáo viên:\n"
            f"--- BẮT ĐẦU MẪU ---\n{sample_content[:4000]}\n--- KẾT THÚC MẪU ---\n\n"
            f"Dựa vào phong cách và cấu trúc của mẫu trên, hãy soạn bài học mới sau:\n"
            f"- Môn học: {subject} - Khối: {grade}\n"
            f"- Tên bài học mới: {title} ({duration})\n"
            f"- Yêu cầu bổ sung: {options_text}\n"
            f"Giữ nguyên văn phong và cách tổ chức các bước của giáo án mẫu."
        )

    response = client.models.generate_content(
        model="gemini-2.5-flash",
        contents=prompt,
        config=types.GenerateContentConfig(
            system_instruction=system_instruction,
            temperature=0.4
        )
    )
    return response.text

def create_docx_file(content, title):
    """Tạo file Word (.docx) chuẩn định dạng in ấn Việt Nam (Times New Roman, căn lề chuẩn)"""
    doc = Document()
    
    # Căn lề chuẩn trang A4 (Trái 3cm, Phải 1.5cm, Trên/Dưới 2cm)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.79)
        section.bottom_margin = Inches(0.79)
        section.left_margin = Inches(1.18)
        section.right_margin = Inches(0.59)

    # Thêm tiêu đề và nội dung
    for line in content.split("\n"):
        line = line.strip()
        if not line:
            continue
            
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(3)
        
        # Nhận diện tiêu đề lớn (I., II., III., IV.)
        if any(line.startswith(prefix) for prefix in ["I.", "II.", "III.", "IV.", "V."]):
            run = p.add_run(line)
            run.font.name = "Times New Roman"
            run.font.size = Pt(13)
            run.font.bold = True
            run.font.color.rgb = RGBColor(30, 58, 138)
        elif line.startswith("#") or line.startswith("KẾ HOẠCH BÀI DẠY") or line.startswith("TÊN BÀI DẠY"):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line.replace("#", "").strip())
            run.font.name = "Times New Roman"
            run.font.size = Pt(14)
            run.font.bold = True
        elif any(line.startswith(prefix) for prefix in ["1.", "2.", "3.", "4.", "a.", "b.", "c.", "d.", "Hoạt động"]):
            run = p.add_run(line)
            run.font.name = "Times New Roman"
            run.font.size = Pt(13)
            run.font.bold = True
        else:
            run = p.add_run(line)
            run.font.name = "Times New Roman"
            run.font.size = Pt(13)

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# ==========================================
# GIAO DIỆN CHÍNH (STREAMLIT UI)
# ==========================================
st.markdown("<div class='main-title'>📚 EduPlan AI - Trợ Lý Soạn Giáo Án</div>", unsafe_allow_html=True)
st.markdown("<div class='sub-title'>Tự động soạn kế hoạch bài dạy chuẩn Công văn 5512 hoặc học theo mẫu Word của trường bạn</div>", unsafe_allow_html=True)

# Chia layout 2 cột: Trái (Bảng điều khiển) - Phải (Vùng xem trước & Tải file)
col_left, col_right = st.columns([1.1, 1.9], gap="medium")

with col_left:
    st.subheader("⚙️ Thiết lập bài soạn")
    
    # Nhập API Key (có thể nhập trực tiếp hoặc qua file secrets)
    gemini_api_key = st.text_input("🔑 Google Gemini API Key:", type="password", help="Lấy miễn phí tại aistudio.google.com")
    
    tab1, tab2 = st.tabs(["⚡ Chuẩn Công văn 5512", "📁 Dựa theo Mẫu Word"])
    
    with tab1:
        st.caption("Sinh giáo án chuẩn mực 4 hoạt động của Bộ GD&ĐT")
        c1, c2 = st.columns(2)
        with c1:
            subject_5512 = st.selectbox("Môn học:", ["Khoa học tự nhiên (Hóa học)", "Khoa học tự nhiên (Sinh học)", "Khoa học tự nhiên (Vật lí)", "Toán học", "Ngữ văn", "Lịch sử & Địa lí", "Khác"])
            grade_5512 = st.selectbox("Khối lớp:", ["Lớp 6", "Lớp 7", "Lớp 8", "Lớp 9"])
        with c2:
            book_5512 = st.selectbox("Bộ sách:", ["Cánh Diều", "Kết Nối Tri Thức", "Chân Trời Sáng Tạo"])
            duration_5512 = st.selectbox("Thời lượng:", ["1 tiết", "2 tiết", "3 tiết", "4 tiết"])
            
        title_5512 = st.text_input("Tên bài học:", value="Bài: Axit - Tính chất và ứng dụng")
        
        st.write("Tùy chọn bổ trợ:")
        opt_ws = st.checkbox("Kèm Phiếu học tập (Worksheet)", value=True)
        opt_quiz = st.checkbox("Kèm 5 câu trắc nghiệm củng cố", value=True)
        opt_game = st.checkbox("Thiết kế trò chơi khởi động / Thí nghiệm", value=True)
        
        options_list_5512 = []
        if opt_ws: options_list_5512.append("Kèm phiếu học tập chi tiết")
        if opt_quiz: options_list_5512.append("Kèm 5 câu hỏi trắc nghiệm đánh giá")
        if opt_game: options_list_5512.append("Thiết kế trò chơi/thí nghiệm khởi động sinh động")
        
        btn_generate_5512 = st.button("🚀 BẮT ĐẦU SOẠN (CHUẨN 5512)", key="btn_5512")

    with tab2:
        st.caption("AI sẽ đọc cấu trúc bảng biểu, font chữ và văn phong của file mẫu để viết bài mới")
        uploaded_file = st.file_uploader("Tải lên file giáo án mẫu (.docx):", type=["docx"])
        
        c1, c2 = st.columns(2)
        with c1:
            subject_custom = st.selectbox("Môn bài mới:", ["Khoa học tự nhiên (Hóa học)", "Khoa học tự nhiên (Sinh học)", "Toán học", "Ngữ văn", "Khác"], key="sb_custom")
            grade_custom = st.selectbox("Khối lớp mới:", ["Lớp 6", "Lớp 7", "Lớp 8", "Lớp 9"], key="gr_custom")
        with c2:
            duration_custom = st.selectbox("Thời lượng bài mới:", ["1 tiết", "2 tiết", "3 tiết"], key="dur_custom")
            book_custom = st.text_input("Bộ sách:", value="Cánh Diều", key="bk_custom")
            
        title_custom = st.text_input("Tên bài học mới cần soạn:", value="Bài: Bazơ - Tính chất hóa học", key="tit_custom")
        
        btn_generate_custom = st.button("🚀 SOẠN BÀI THEO MẪU ĐÃ TẢI", key="btn_custom")

# ==========================================
# XỬ LÝ SỰ KIỆN NÚT BẤM
# ==========================================
if btn_generate_5512:
    if not gemini_api_key:
        st.error("⚠️ Vui lòng nhập Google Gemini API Key ở trên để tiếp tục!")
    elif not title_5512:
        st.warning("⚠️ Vui lòng nhập tên bài học!")
    else:
        with st.spinner("⏳ Đang tham chiếu chuẩn GDPT 2018 và sinh giáo án chi tiết..."):
            try:
                res = generate_lesson_plan(
                    api_key=gemini_api_key,
                    mode="5512",
                    subject=subject_5512,
                    grade=grade_5512,
                    textbook=book_5512,
                    title=title_5512,
                    duration=duration_5512,
                    options=options_list_5512
                )
                st.session_state["lesson_content"] = res
                st.session_state["lesson_title"] = f"GiaoAn_{grade_5512}_{title_5512.replace(' ', '_')}"
                st.success("✅ Đã tạo xong giáo án thành công!")
            except Exception as e:
                st.error(f"❌ Có lỗi xảy ra: {str(e)}")

if btn_generate_custom:
    if not gemini_api_key:
        st.error("⚠️ Vui lòng nhập Google Gemini API Key ở trên để tiếp tục!")
    elif not uploaded_file:
        st.warning("⚠️ Vui lòng chọn và tải lên 1 file Word giáo án mẫu (.docx)!")
    elif not title_custom:
        st.warning("⚠️ Vui lòng nhập tên bài học mới cần soạn!")
    else:
        with st.spinner("⏳ Đang phân tích phong cách file mẫu và soạn bài mới..."):
            try:
                sample_text = extract_text_from_docx(uploaded_file)
                res = generate_lesson_plan(
                    api_key=gemini_api_key,
                    mode="custom",
                    subject=subject_custom,
                    grade=grade_custom,
                    textbook=book_custom,
                    title=title_custom,
                    duration=duration_custom,
                    options=["Bắt chước phong cách bài mẫu"],
                    sample_content=sample_text
                )
                st.session_state["lesson_content"] = res
                st.session_state["lesson_title"] = f"GiaoAn_TheoMau_{title_custom.replace(' ', '_')}"
                st.success("✅ Đã hoàn thành bài soạn theo đúng mẫu của bạn!")
            except Exception as e:
                st.error(f"❌ Có lỗi xảy ra: {str(e)}")

# ==========================================
# CỘT PHẢI: KHUNG XEM TRƯỚC VÀ XUẤT FILE WORD
# ==========================================
with col_right:
    st.subheader("📄 Bản xem trước & Chỉnh sửa trực tiếp")
    
    if st.session_state["lesson_content"]:
        # Khung cho phép giáo viên chỉnh sửa trực tiếp nội dung trước khi xuất
        edited_text = st.text_area(
            "Bạn có thể gõ sửa trực tiếp vào khung dưới đây trước khi tải về:",
            value=st.session_state["lesson_content"],
            height=550
        )
        
        # Cập nhật lại nội dung nếu có sửa
        st.session_state["lesson_content"] = edited_text
        
        # Nút tạo và tải file Word
        docx_file = create_docx_file(st.session_state["lesson_content"], st.session_state["lesson_title"])
        
        st.download_button(
            label="📥 TẢI XUỐNG FILE WORD (.DOCX) CHUẨN IN ẤN",
            data=docx_file,
            file_name=f"{st.session_state['lesson_title']}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary"
        )
    else:
        st.info("👈 Hãy chọn thông tin hoặc tải file mẫu ở cột bên trái, sau đó bấm nút **Bắt đầu soạn** để xem kết quả tại đây.")
